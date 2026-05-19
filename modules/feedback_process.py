"""
modules/feedback_process.py
Sistema de Feedback Process replicando el template DOCX corporativo.

Sofi llena el form → se guarda → genera link público → empleado firma desde WhatsApp.
Permite descargar PDF y DOCX del feedback.
"""
import streamlit as st
import pandas as pd
import uuid
from datetime import date, datetime, timedelta
from io import BytesIO

from core.ui import render_page_title
from core.sheets import read_worksheet, append_row, get_worksheet, invalidate_cache
from core.config import WS_FEEDBACK_PROCESS, WS_EMPLOYEES
from core.time_utils import today_gt, now_gt, parse_date
from core.flags import flag_emoji_unicode
from core.auth import current_user_display_name
from core.notifications import notify_success, notify_error


TIPOS_FEEDBACK = [
    "Positive / Recognition",
    "Corrective",
    "Preventive",
    "Informational",
]

AREAS_FEEDBACK = [
    "Performance / Task Execution",
    "Behavior / Professional Conduct",
    "Communication",
    "Process Adherence",
    "Attendance / Punctuality",
    "Teamwork / Collaboration",
    "Other",
]


def render():
    render_page_title(
        eyebrow="RECURSOS HUMANOS",
        title="Feedback Process",
        subtitle="Documenta feedback individual con firma virtual del empleado",
    )

    try:
        employees_df = read_worksheet(WS_EMPLOYEES)
    except Exception as e:
        st.error(f"Error: {e}")
        return

    employees_active = employees_df[
        employees_df["activo"].astype(str).str.upper().isin(["TRUE", "VERDADERO", "SI", "1"])
    ].copy() if not employees_df.empty else pd.DataFrame()

    tab_new, tab_history = st.tabs([
        "➕  Nuevo feedback",
        "📜  Historial",
    ])

    with tab_new:
        _render_new_form(employees_active)

    with tab_history:
        _render_history(employees_active)


def _render_new_form(employees_active: pd.DataFrame):
    """Form replicando el template DOCX."""

    st.markdown(
        '<div style="font-size:13px;color:#64748B;margin-bottom:16px;'
        'padding:12px 16px;background:#F8FAFC;border-left:3px solid #2563EB;'
        'border-radius:0 4px 4px 0;">'
        '<strong style="color:#0A0A0A">Feedback Process</strong><br>'
        'Documenta el feedback de forma específica, oportuna y constructiva. '
        'Al guardar se genera un link único para compartir con el empleado por WhatsApp, '
        'donde podrá leer y firmar virtualmente.'
        '</div>',
        unsafe_allow_html=True,
    )

    # ============================================================
    # SECCIÓN: General Information
    # ============================================================
    st.markdown(
        '<div style="font-size:11px;font-weight:700;letter-spacing:1.5px;'
        'text-transform:uppercase;color:#DC2626;margin:16px 0 12px 0;">'
        '— GENERAL INFORMATION'
        '</div>',
        unsafe_allow_html=True,
    )

    emp_options = {"": None}
    for _, emp in employees_active.iterrows():
        flag_uc = flag_emoji_unicode(emp.get("pais", ""))
        display = f"{flag_uc}  {emp['nombre']} ({emp.get('rol', '')})"
        emp_options[display] = {
            "id": int(emp["id"]),
            "nombre": emp["nombre"],
            "rol": emp.get("rol", ""),
        }

    col1, col2 = st.columns(2)
    with col1:
        emp_display = st.selectbox(
            "Employee Name",
            options=list(emp_options.keys()),
            key="fp_emp",
        )
    with col2:
        fecha_fb = st.date_input(
            "Date of Feedback",
            value=today_gt(),
            format="DD/MM/YYYY",
            key="fp_fecha",
        )

    selected_emp = emp_options.get(emp_display)

    col3, col4 = st.columns(2)
    with col3:
        posicion = st.text_input(
            "Position",
            value=selected_emp["rol"] if selected_emp else "",
            key="fp_pos",
            max_chars=80,
        )
    with col4:
        departamento = st.text_input(
            "Department",
            value="BackOffice",
            key="fp_dept",
            max_chars=80,
        )

    manager = st.text_input(
        "Manager / Supervisor Providing Feedback",
        value=current_user_display_name(),
        key="fp_manager",
        max_chars=80,
    )

    # ============================================================
    # SECCIÓN: Type of Feedback
    # ============================================================
    st.markdown(
        '<div style="font-size:11px;font-weight:700;letter-spacing:1.5px;'
        'text-transform:uppercase;color:#DC2626;margin:24px 0 12px 0;">'
        '— TYPE OF FEEDBACK'
        '</div>',
        unsafe_allow_html=True,
    )

    tipo_feedback = st.radio(
        "Type",
        options=TIPOS_FEEDBACK,
        horizontal=True,
        key="fp_tipo",
        label_visibility="collapsed",
    )

    # ============================================================
    # SECCIÓN: Feedback Area
    # ============================================================
    st.markdown(
        '<div style="font-size:11px;font-weight:700;letter-spacing:1.5px;'
        'text-transform:uppercase;color:#DC2626;margin:24px 0 12px 0;">'
        '— FEEDBACK AREA'
        '</div>',
        unsafe_allow_html=True,
    )

    area_feedback = st.selectbox(
        "Select the primary area",
        options=AREAS_FEEDBACK,
        key="fp_area",
        label_visibility="collapsed",
    )

    area_otro = ""
    if area_feedback == "Other":
        area_otro = st.text_input("Specify other area", key="fp_area_otro", max_chars=100)

    # ============================================================
    # SECCIÓN: Description of the Situation
    # ============================================================
    st.markdown(
        '<div style="font-size:11px;font-weight:700;letter-spacing:1.5px;'
        'text-transform:uppercase;color:#DC2626;margin:24px 0 12px 0;">'
        '— DESCRIPTION OF THE SITUATION'
        '</div>',
        unsafe_allow_html=True,
    )
    descripcion_situacion = st.text_area(
        "Describe objectively: who, what, when, where. Avoid assumptions.",
        height=120,
        key="fp_desc",
        placeholder="Ej: On October 15 during the morning shift, the employee...",
        max_chars=2000,
    )

    # ============================================================
    # SECCIÓN: Feedback Provided
    # ============================================================
    st.markdown(
        '<div style="font-size:11px;font-weight:700;letter-spacing:1.5px;'
        'text-transform:uppercase;color:#DC2626;margin:24px 0 12px 0;">'
        '— FEEDBACK PROVIDED'
        '</div>',
        unsafe_allow_html=True,
    )
    feedback_dado = st.text_area(
        "Clearly state the feedback given to the employee. Be specific and constructive.",
        height=120,
        key="fp_fb_dado",
        max_chars=2000,
    )

    # ============================================================
    # SECCIÓN: Expected Behavior or Improvement
    # ============================================================
    st.markdown(
        '<div style="font-size:11px;font-weight:700;letter-spacing:1.5px;'
        'text-transform:uppercase;color:#DC2626;margin:24px 0 12px 0;">'
        '— EXPECTED BEHAVIOR OR IMPROVEMENT'
        '</div>',
        unsafe_allow_html=True,
    )
    comportamiento_esperado = st.text_area(
        "Explain what is expected moving forward.",
        height=100,
        key="fp_exp",
        max_chars=1500,
    )

    # ============================================================
    # SECCIÓN: Action Items
    # ============================================================
    st.markdown(
        '<div style="font-size:11px;font-weight:700;letter-spacing:1.5px;'
        'text-transform:uppercase;color:#DC2626;margin:24px 0 12px 0;">'
        '— ACTION ITEMS (If Applicable)'
        '</div>',
        unsafe_allow_html=True,
    )
    accion_empleado = st.text_area(
        "Action to be taken by the employee",
        height=80, key="fp_acc_emp", max_chars=800,
    )
    apoyo_manager = st.text_area(
        "Support or guidance provided by the manager",
        height=80, key="fp_apoyo", max_chars=800,
    )
    fecha_seguimiento = st.date_input(
        "Deadline or follow-up date (if applicable)",
        value=today_gt() + timedelta(days=14),
        format="DD/MM/YYYY",
        key="fp_seg",
    )

    # ============================================================
    # SECCIÓN: Follow-Up
    # ============================================================
    st.markdown(
        '<div style="font-size:11px;font-weight:700;letter-spacing:1.5px;'
        'text-transform:uppercase;color:#DC2626;margin:24px 0 12px 0;">'
        '— FOLLOW-UP REQUIRED'
        '</div>',
        unsafe_allow_html=True,
    )
    followup_required = st.radio(
        "Follow-up required?",
        options=["No", "Yes"],
        horizontal=True,
        key="fp_fu_req",
    )
    followup_date = None
    followup_notes = ""
    if followup_required == "Yes":
        followup_date = st.date_input(
            "Follow-up date",
            value=today_gt() + timedelta(days=30),
            format="DD/MM/YYYY",
            key="fp_fu_date",
        )
        followup_notes = st.text_area("Additional notes", height=80, key="fp_fu_notes")

    # ============================================================
    # BOTÓN: GUARDAR
    # ============================================================
    st.divider()

    if st.button(
        "💾 Guardar feedback y generar link de firma",
        use_container_width=True,
        type="primary",
        key="fp_save",
        disabled=(selected_emp is None or not feedback_dado.strip() or not descripcion_situacion.strip()),
    ):
        try:
            fb_id = f"FB-{now_gt().strftime('%Y%m%d%H%M%S')}-{uuid.uuid4().hex[:8].upper()}"
            timestamp = now_gt().strftime("%Y-%m-%d %H:%M:%S")

            row = [
                fb_id,
                fecha_fb.strftime("%Y-%m-%d"),
                selected_emp["id"],
                selected_emp["nombre"],
                posicion,
                departamento,
                manager,
                tipo_feedback,
                area_feedback,
                area_otro,
                descripcion_situacion,
                feedback_dado,
                comportamiento_esperado,
                accion_empleado,
                apoyo_manager,
                fecha_seguimiento.strftime("%Y-%m-%d") if fecha_seguimiento else "",
                "", "",  # empleado_acknowledged, comentario_empleado (vacíos hasta que firme)
                followup_required,
                followup_date.strftime("%Y-%m-%d") if followup_date else "",
                followup_notes,
                "PENDIENTE_FIRMA",
                "",  # fecha_firma
                "",  # comentario_firma
                "",  # ip_firma
                timestamp,
                timestamp,
            ]
            append_row(WS_FEEDBACK_PROCESS, row)
            invalidate_cache()
            notify_success(
                f"Feedback para {selected_emp['nombre']} guardado. Ve al historial para compartir el link.",
                title="Feedback guardado"
            )
            st.session_state["last_saved_fb_id"] = fb_id
            st.rerun()
        except Exception as e:
            notify_error(str(e))

    # Si acabamos de guardar, mostrar info del link
    last_fb = st.session_state.get("last_saved_fb_id")
    if last_fb:
        share_link = _build_share_link(last_fb)
        st.success(
            f"✅ Feedback guardado con ID **{last_fb}**.\n\n"
            f"**Link para enviar al empleado:**\n\n"
            f"```\n{share_link}\n```\n\n"
            f"Cópialo y pégalo en WhatsApp, o usa el botón en el historial."
        )


def _render_history(employees_active: pd.DataFrame):
    """Histórico de feedbacks enviados."""
    try:
        df = read_worksheet(WS_FEEDBACK_PROCESS)
    except Exception as e:
        st.error(f"Error: {e}")
        return

    if df.empty:
        st.info("📭 No hay feedbacks registrados todavía.")
        return

    df["fecha_parsed"] = df["fecha"].apply(parse_date)
    df = df[df["fecha_parsed"].notna()].sort_values("fecha_parsed", ascending=False).reset_index(drop=True)

    # KPIs
    total = len(df)
    firmados = int((df["estado_firma"].astype(str).str.upper() == "FIRMADO").sum())
    pendientes = total - firmados

    kpi_html = f"""
    <div style="display:grid;grid-template-columns:repeat(3,1fr);gap:12px;margin:16px 0 20px 0;">
        <div style="background:#FFFFFF;border:1px solid #E2E8F0;border-radius:6px;padding:16px 18px;">
            <div style="font-size:10px;font-weight:700;letter-spacing:1.5px;text-transform:uppercase;color:#64748B;margin-bottom:8px;">TOTAL FEEDBACKS</div>
            <div style="font-size:32px;font-weight:700;color:#0A0A0A;line-height:1;letter-spacing:-1px;">{total}</div>
        </div>
        <div style="background:#FFFFFF;border:1px solid #E2E8F0;border-radius:6px;padding:16px 18px;">
            <div style="font-size:10px;font-weight:700;letter-spacing:1.5px;text-transform:uppercase;color:#64748B;margin-bottom:8px;">FIRMADOS</div>
            <div style="font-size:32px;font-weight:700;color:#16A34A;line-height:1;letter-spacing:-1px;">{firmados}</div>
        </div>
        <div style="background:#FFFFFF;border:1px solid #E2E8F0;border-radius:6px;padding:16px 18px;">
            <div style="font-size:10px;font-weight:700;letter-spacing:1.5px;text-transform:uppercase;color:#64748B;margin-bottom:8px;">PENDIENTES</div>
            <div style="font-size:32px;font-weight:700;color:#D97706;line-height:1;letter-spacing:-1px;">{pendientes}</div>
        </div>
    </div>
    """
    st.markdown(kpi_html, unsafe_allow_html=True)

    for _, row in df.iterrows():
        fb_id = row.get("id", "")
        fecha_str = row["fecha_parsed"].strftime("%d/%m/%Y") if row["fecha_parsed"] else ""
        emp_name = row.get("empleado_nombre", "")
        tipo = row.get("tipo_feedback", "")
        estado = str(row.get("estado_firma", "")).upper()

        if estado == "FIRMADO":
            estado_label = "✓ FIRMADO"
            estado_color = "#16A34A"
        else:
            estado_label = "⏳ PENDIENTE"
            estado_color = "#D97706"

        with st.expander(f"{estado_label}  ·  {emp_name}  ·  {tipo}  ·  {fecha_str}"):
            st.markdown(f"**Manager:** {row.get('manager', '')}")
            st.markdown(f"**Área:** {row.get('area_feedback', '')}")
            st.markdown(f"**Descripción:** {row.get('descripcion_situacion', '')}")
            st.markdown(f"**Feedback dado:** {row.get('feedback_dado', '')}")
            st.markdown(f"**Comportamiento esperado:** {row.get('comportamiento_esperado', '')}")

            if estado == "FIRMADO":
                st.success(
                    f"Firmado el {row.get('fecha_firma', '')}\n\n"
                    f"**Comentario del empleado:** {row.get('comentario_firma', '') or '—'}"
                )

            share_link = _build_share_link(fb_id)
            st.markdown(f"**Link para WhatsApp:**\n\n```\n{share_link}\n```")

            wa_link = _build_whatsapp_link(share_link, emp_name)

            col_pdf, col_docx, col_wa = st.columns(3)
            with col_pdf:
                pdf_buffer = _generate_feedback_pdf(row)
                if pdf_buffer:
                    st.download_button(
                        "📄 PDF",
                        data=pdf_buffer,
                        file_name=f"Feedback_{emp_name}_{fecha_str.replace('/', '-')}.pdf",
                        mime="application/pdf",
                        key=f"pdf_{fb_id}",
                        use_container_width=True,
                    )
            with col_docx:
                docx_buffer = _generate_feedback_docx(row)
                if docx_buffer:
                    st.download_button(
                        "📝 DOCX",
                        data=docx_buffer,
                        file_name=f"Feedback_{emp_name}_{fecha_str.replace('/', '-')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=f"docx_{fb_id}",
                        use_container_width=True,
                    )
            with col_wa:
                st.link_button(
                    "💬 WhatsApp",
                    wa_link,
                    use_container_width=True,
                )


def _build_share_link(fb_id: str) -> str:
    """Construye el link público para que el empleado firme."""
    from core.config import APP_URL
    base = APP_URL.rstrip("/") if APP_URL else "https://attendance-backoffice-stt.streamlit.app"
    return f"{base}/?feedback={fb_id}"


def _build_whatsapp_link(share_link: str, emp_name: str) -> str:
    """Construye un wa.me link prellenado."""
    import urllib.parse
    msg = (
        f"Hola {emp_name}! Te comparto un feedback formal que necesito que revises y firmes virtualmente.\n\n"
        f"Abre este link:\n{share_link}\n\n"
        f"Gracias!"
    )
    return f"https://wa.me/?text={urllib.parse.quote(msg)}"


def _generate_feedback_pdf(row):
    """Genera PDF del feedback con formato corporativo STT."""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib.colors import HexColor
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        )
        from reportlab.lib.enums import TA_LEFT, TA_CENTER
    except ImportError:
        return None

    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=letter,
        rightMargin=0.6*inch, leftMargin=0.6*inch,
        topMargin=0.5*inch, bottomMargin=0.5*inch,
    )

    styles = getSampleStyleSheet()
    story = []

    header_style = ParagraphStyle(
        "h", parent=styles["Normal"], fontSize=20, textColor=HexColor("#DC2626"),
        spaceAfter=4, fontName="Helvetica-Bold",
    )
    subtitle_style = ParagraphStyle(
        "sub", parent=styles["Normal"], fontSize=9, textColor=HexColor("#64748B"),
        spaceAfter=14,
    )
    section_style = ParagraphStyle(
        "sec", parent=styles["Normal"], fontSize=10, textColor=HexColor("#DC2626"),
        spaceBefore=12, spaceAfter=6, fontName="Helvetica-Bold",
    )
    body_style = ParagraphStyle(
        "body", parent=styles["Normal"], fontSize=9, textColor=HexColor("#0A0A0A"),
        spaceAfter=4, leading=12,
    )

    story.append(Paragraph("FEEDBACK PROCESS", header_style))
    story.append(Paragraph(
        "Confidential HR documentation · STT Logistics Group",
        subtitle_style,
    ))

    # General Information
    story.append(Paragraph("— GENERAL INFORMATION", section_style))
    info_data = [
        ["Employee Name:", _esc(row.get("empleado_nombre"))],
        ["Position:", _esc(row.get("posicion"))],
        ["Department:", _esc(row.get("departamento"))],
        ["Manager/Supervisor:", _esc(row.get("manager"))],
        ["Date of Feedback:", _esc(row.get("fecha"))],
    ]
    t = Table(info_data, colWidths=[2*inch, 4.5*inch])
    t.setStyle(TableStyle([
        ("FONTSIZE", (0,0), (-1,-1), 9),
        ("FONTNAME", (0,0), (0,-1), "Helvetica-Bold"),
        ("TEXTCOLOR", (0,0), (0,-1), HexColor("#64748B")),
        ("BOTTOMPADDING", (0,0), (-1,-1), 6),
    ]))
    story.append(t)

    # Type & Area
    story.append(Paragraph("— TYPE OF FEEDBACK", section_style))
    story.append(Paragraph(_esc(row.get("tipo_feedback")), body_style))

    story.append(Paragraph("— FEEDBACK AREA", section_style))
    area = row.get("area_feedback", "")
    if area == "Other" and row.get("area_otro"):
        area = f"Other: {row.get('area_otro')}"
    story.append(Paragraph(_esc(area), body_style))

    # Description
    story.append(Paragraph("— DESCRIPTION OF THE SITUATION", section_style))
    story.append(Paragraph(_esc(row.get("descripcion_situacion")), body_style))

    # Feedback Provided
    story.append(Paragraph("— FEEDBACK PROVIDED", section_style))
    story.append(Paragraph(_esc(row.get("feedback_dado")), body_style))

    # Expected
    story.append(Paragraph("— EXPECTED BEHAVIOR OR IMPROVEMENT", section_style))
    story.append(Paragraph(_esc(row.get("comportamiento_esperado")), body_style))

    # Action Items
    if row.get("accion_empleado") or row.get("apoyo_manager"):
        story.append(Paragraph("— ACTION ITEMS", section_style))
        if row.get("accion_empleado"):
            story.append(Paragraph(f"<b>Employee action:</b> {_esc(row.get('accion_empleado'))}", body_style))
        if row.get("apoyo_manager"):
            story.append(Paragraph(f"<b>Manager support:</b> {_esc(row.get('apoyo_manager'))}", body_style))
        if row.get("fecha_seguimiento"):
            story.append(Paragraph(f"<b>Follow-up date:</b> {_esc(row.get('fecha_seguimiento'))}", body_style))

    # Acknowledgment
    estado = str(row.get("estado_firma", "")).upper()
    story.append(Paragraph("— EMPLOYEE ACKNOWLEDGMENT", section_style))
    if estado == "FIRMADO":
        story.append(Paragraph(
            f"<b>Acknowledged:</b> ☑ Yes · Signed virtually on {_esc(row.get('fecha_firma'))}",
            body_style,
        ))
        if row.get("comentario_firma"):
            story.append(Paragraph(f"<b>Employee comment:</b> {_esc(row.get('comentario_firma'))}", body_style))
    else:
        story.append(Paragraph("<b>Acknowledged:</b> ☐ Pending signature", body_style))

    # Signatures
    story.append(Spacer(1, 0.25*inch))
    story.append(Paragraph("This feedback session has been discussed with the employee.", body_style))
    story.append(Spacer(1, 0.2*inch))

    sig_data = [
        ["Employee Signature: _________________", "Date: _________"],
        ["Manager Signature:  _________________", "Date: _________"],
        ["HR:                  _________________", "Date: _________"],
    ]
    sig_t = Table(sig_data, colWidths=[4*inch, 2.5*inch])
    sig_t.setStyle(TableStyle([
        ("FONTSIZE", (0,0), (-1,-1), 9),
        ("BOTTOMPADDING", (0,0), (-1,-1), 10),
    ]))
    story.append(sig_t)

    # Footer
    story.append(Spacer(1, 0.3*inch))
    footer_style = ParagraphStyle(
        "ft", parent=styles["Normal"], fontSize=7, textColor=HexColor("#94A3B8"),
        alignment=TA_CENTER,
    )
    story.append(Paragraph("Confidential and proprietary. © STT Logistics Group. All Rights Reserved.", footer_style))

    doc.build(story)
    buffer.seek(0)
    return buffer


def _generate_feedback_docx(row):
    """Genera DOCX del feedback."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return None

    doc = Document()

    # Título
    title = doc.add_paragraph()
    run = title.add_run("FEEDBACK PROCESS")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)

    sub = doc.add_paragraph()
    sub_run = sub.add_run("Confidential HR documentation · STT Logistics Group")
    sub_run.font.size = Pt(9)
    sub_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph()

    def _add_section(label, content):
        if not content:
            return
        h = doc.add_paragraph()
        hr = h.add_run(label)
        hr.bold = True
        hr.font.size = Pt(11)
        hr.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
        p = doc.add_paragraph(str(content))
        for r in p.runs:
            r.font.size = Pt(10)

    _add_section("General Information", "")
    info_table = doc.add_table(rows=5, cols=2)
    info_table.style = "Light Grid Accent 1"
    info_table.cell(0, 0).text = "Employee Name"
    info_table.cell(0, 1).text = str(row.get("empleado_nombre") or "")
    info_table.cell(1, 0).text = "Position"
    info_table.cell(1, 1).text = str(row.get("posicion") or "")
    info_table.cell(2, 0).text = "Department"
    info_table.cell(2, 1).text = str(row.get("departamento") or "")
    info_table.cell(3, 0).text = "Manager/Supervisor"
    info_table.cell(3, 1).text = str(row.get("manager") or "")
    info_table.cell(4, 0).text = "Date of Feedback"
    info_table.cell(4, 1).text = str(row.get("fecha") or "")

    _add_section("Type of Feedback", row.get("tipo_feedback"))
    area = row.get("area_feedback", "")
    if area == "Other" and row.get("area_otro"):
        area = f"Other: {row.get('area_otro')}"
    _add_section("Feedback Area", area)
    _add_section("Description of the Situation", row.get("descripcion_situacion"))
    _add_section("Feedback Provided", row.get("feedback_dado"))
    _add_section("Expected Behavior or Improvement", row.get("comportamiento_esperado"))

    if row.get("accion_empleado") or row.get("apoyo_manager"):
        h = doc.add_paragraph()
        hr = h.add_run("Action Items")
        hr.bold = True
        hr.font.size = Pt(11)
        hr.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
        if row.get("accion_empleado"):
            doc.add_paragraph(f"Employee action: {row.get('accion_empleado')}")
        if row.get("apoyo_manager"):
            doc.add_paragraph(f"Manager support: {row.get('apoyo_manager')}")
        if row.get("fecha_seguimiento"):
            doc.add_paragraph(f"Follow-up date: {row.get('fecha_seguimiento')}")

    # Acknowledgment
    estado = str(row.get("estado_firma", "")).upper()
    _add_section("Employee Acknowledgment",
                 f"Signed virtually on {row.get('fecha_firma')}" if estado == "FIRMADO" else "Pending signature")
    if estado == "FIRMADO" and row.get("comentario_firma"):
        doc.add_paragraph(f"Employee comment: {row.get('comentario_firma')}")

    # Signatures
    doc.add_paragraph()
    doc.add_paragraph("This feedback session has been discussed with the employee.")
    doc.add_paragraph()
    doc.add_paragraph("Employee Signature: __________________   Date: __________")
    doc.add_paragraph("Manager Signature:  __________________   Date: __________")
    doc.add_paragraph("HR:                  __________________   Date: __________")

    doc.add_paragraph()
    ft = doc.add_paragraph()
    ft.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ftr = ft.add_run("Confidential and proprietary. © STT Logistics Group. All Rights Reserved.")
    ftr.font.size = Pt(8)
    ftr.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _esc(text):
    if text is None:
        return ""
    s = str(text)
    return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
